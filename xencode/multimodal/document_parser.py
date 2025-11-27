"""Document parser for PDF and DOCX files."""

from pathlib import Path
from typing import Dict, Any, List
import re

try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DocumentParser:
    """Parser for document files (PDF, DOCX)."""

    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document file to extract text and metadata.
        
        Args:
            file_path: Path to the document file.
            
        Returns:
            Dictionary containing text content and metadata.
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Document file not found: {file_path}")

        result = {
            "filename": path.name,
            "path": str(path),
            "file_type": path.suffix.lower(),
            "text": "",
            "metadata": {},
            "error": None
        }

        try:
            if path.suffix.lower() == ".pdf":
                result.update(self._parse_pdf(path))
            elif path.suffix.lower() == ".docx":
                result.update(self._parse_docx(path))
            else:
                result["error"] = f"Unsupported file type: {path.suffix}"
        except Exception as e:
            result["error"] = str(e)

        return result

    def _parse_pdf(self, path: Path) -> Dict[str, Any]:
        """Parse PDF file."""
        if not PYPDF_AVAILABLE:
            return {"error": "pypdf is not installed. Run: pip install pypdf"}

        reader = PdfReader(str(path))
        
        text_parts = []
        for page in reader.pages:
            text_parts.append(page.extract_text())
        
        metadata = {
            "num_pages": len(reader.pages),
            "pdf_metadata": {}
        }
        
        if reader.metadata:
            metadata["pdf_metadata"] = {
                k: str(v) for k, v in reader.metadata.items() if v
            }

        return {
            "text": "\n\n".join(text_parts),
            "metadata": metadata
        }

    def _parse_docx(self, path: Path) -> Dict[str, Any]:
        """Parse DOCX file."""
        if not DOCX_AVAILABLE:
            return {"error": "python-docx is not installed. Run: pip install python-docx"}

        doc = Document(str(path))
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        metadata = {
            "num_paragraphs": len(doc.paragraphs),
            "num_tables": len(doc.tables),
            "num_sections": len(doc.sections)
        }
        
        # Extract core properties if available
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            metadata["docx_metadata"] = {
                "author": props.author or "",
                "title": props.title or "",
                "subject": props.subject or "",
                "created": str(props.created) if props.created else "",
                "modified": str(props.modified) if props.modified else ""
            }

        return {
            "text": "\n\n".join(text_parts),
            "metadata": metadata
        }

    def clean_text(self, text: str) -> str:
        """Clean extracted text by removing excess whitespace."""
        # Replace multiple newlines with double newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Replace multiple spaces with single space
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()
