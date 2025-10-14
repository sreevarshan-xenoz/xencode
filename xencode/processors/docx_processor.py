#!/usr/bin/env python3
"""
DOCX Document Processor

Implements DOCX processing using python-docx for text extraction,
table processing, metadata extraction, and structured content analysis.
"""

import asyncio
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    DocxDocument = None

from xencode.processors.base import ProcessorInterface
from xencode.models.document import (
    ContentType,
    DocumentMetadata,
    DocumentType,
    ProcessedDocument,
    ProcessingOptions,
    ProcessingStatus,
    StructuredContent
)


class DOCXProcessor(ProcessorInterface):
    """DOCX document processor using python-docx"""
    
    def __init__(self):
        if not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is required for DOCX processing. "
                "Install with: pip install python-docx"
            )
        
        self.supported_types = [DocumentType.DOCX]
    
    async def can_process(self, file_path: Path, document_type: DocumentType) -> bool:
        """Check if this processor can handle the document"""
        if document_type != DocumentType.DOCX:
            return False
        
        if not PYTHON_DOCX_AVAILABLE:
            return False
        
        try:
            # Try to open the DOCX to verify it's valid
            doc = DocxDocument(str(file_path))
            return True
        except Exception:
            return False
    
    def get_supported_types(self) -> List[DocumentType]:
        """Get supported document types"""
        return self.supported_types.copy()
    
    async def process(self, 
                     file_path: Path, 
                     options: ProcessingOptions) -> ProcessedDocument:
        """Process DOCX document and extract structured content"""
        
        document = ProcessedDocument(
            original_filename=file_path.name,
            file_path=str(file_path),
            document_type=DocumentType.DOCX,
            processing_status=ProcessingStatus.PROCESSING
        )
        
        try:
            # Open DOCX document
            docx_doc = DocxDocument(str(file_path))
            
            # Extract metadata
            if options.extract_metadata:
                document.metadata = await self._extract_metadata(docx_doc, file_path)
            
            # Extract text and structured content
            if options.extract_text or options.extract_structured_content:
                await self._extract_content(docx_doc, document, options)
            
            # Extract tables
            if options.include_tables:
                await self._extract_tables(docx_doc, document)
            
            # Calculate confidence score
            document.confidence_score = self._calculate_confidence_score(document)
            
            # Set processing status
            document.processing_status = ProcessingStatus.COMPLETED
            
        except Exception as e:
            document.processing_status = ProcessingStatus.FAILED
            document.errors.append(f"DOCX processing error: {str(e)}")
            document.confidence_score = 0.0
        
        return document
    
    async def _extract_metadata(self, 
                               docx_doc: 'DocxDocument', 
                               file_path: Path) -> DocumentMetadata:
        """Extract DOCX metadata"""
        metadata = DocumentMetadata()
        
        try:
            # Get document properties
            core_props = docx_doc.core_properties
            
            metadata.title = core_props.title
            metadata.author = core_props.author
            metadata.created_date = core_props.created
            metadata.modified_date = core_props.modified
            metadata.file_size = file_path.stat().st_size
            metadata.mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            
            # DOCX-specific metadata
            metadata.docx_version = "DOCX"
            
            # Check for macros (VBA)
            try:
                # This is a basic check - more sophisticated detection would be needed
                metadata.has_macros = False  # python-docx doesn't easily expose VBA info
            except:
                metadata.has_macros = False
            
            # Count paragraphs and estimate word count
            paragraph_count = len(docx_doc.paragraphs)
            word_count = 0
            
            for paragraph in docx_doc.paragraphs:
                if paragraph.text.strip():
                    word_count += len(paragraph.text.split())
            
            # Add words from tables
            for table in docx_doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            word_count += len(cell.text.split())
            
            metadata.word_count = word_count
            
        except Exception as e:
            # Continue processing even if metadata extraction fails
            pass
        
        return metadata
    
    async def _extract_content(self, 
                              docx_doc: 'DocxDocument',
                              document: ProcessedDocument,
                              options: ProcessingOptions) -> None:
        """Extract text and structured content from DOCX"""
        
        all_text_parts = []
        
        # Process paragraphs
        for para_index, paragraph in enumerate(docx_doc.paragraphs):
            para_text = paragraph.text.strip()
            
            if not para_text:
                continue
            
            # Add to full text
            if options.extract_text:
                all_text_parts.append(para_text)
            
            # Extract structured content
            if options.extract_structured_content:
                await self._extract_structured_content_from_paragraph(
                    paragraph, para_index, document, options
                )
        
        # Set extracted text
        if options.extract_text:
            document.extracted_text = '\n\n'.join(all_text_parts)
    
    async def _extract_structured_content_from_paragraph(self,
                                                        paragraph,
                                                        para_index: int,
                                                        document: ProcessedDocument,
                                                        options: ProcessingOptions) -> None:
        """Extract structured content from a paragraph"""
        
        try:
            para_text = paragraph.text.strip()
            if not para_text:
                return
            
            # Get paragraph formatting
            formatting = self._get_paragraph_formatting(paragraph)
            
            # Determine content type
            content_type = self._determine_content_type(para_text, formatting, paragraph)
            
            # Create structured content
            structured_content = StructuredContent(
                content_type=content_type,
                text=para_text,
                position={
                    'paragraph_index': para_index
                },
                formatting=formatting,
                confidence=0.9
            )
            
            document.structured_content.append(structured_content)
            
        except Exception as e:
            document.warnings.append(f"Error extracting structured content from paragraph {para_index}: {str(e)}")
    
    def _get_paragraph_formatting(self, paragraph) -> Dict[str, Any]:
        """Extract formatting information from paragraph"""
        formatting = {}
        
        try:
            # Style information
            if paragraph.style:
                formatting['style_name'] = paragraph.style.name
                formatting['style_type'] = str(paragraph.style.type)
            
            # Alignment
            if paragraph.alignment:
                formatting['alignment'] = str(paragraph.alignment)
            
            # Run-level formatting (from first run)
            if paragraph.runs:
                first_run = paragraph.runs[0]
                
                formatting['bold'] = first_run.bold
                formatting['italic'] = first_run.italic
                formatting['underline'] = first_run.underline
                
                if first_run.font.size:
                    formatting['font_size'] = first_run.font.size.pt
                
                if first_run.font.name:
                    formatting['font_name'] = first_run.font.name
        
        except Exception:
            # Continue if formatting extraction fails
            pass
        
        return formatting
    
    def _determine_content_type(self, 
                               text: str, 
                               formatting: Dict[str, Any],
                               paragraph) -> ContentType:
        """Determine content type based on text, formatting, and style"""
        
        # Check style-based content types
        style_name = formatting.get('style_name', '').lower()
        
        # Heading detection
        if 'heading' in style_name or 'title' in style_name:
            return ContentType.HEADING
        
        # Check formatting-based detection
        is_bold = formatting.get('bold', False)
        font_size = formatting.get('font_size', 0)
        
        # Heading detection based on formatting
        if is_bold and font_size > 12:
            if (len(text) < 100 and 
                (text.isupper() or 
                 re.match(r'^\d+\.?\s+[A-Z]', text) or
                 re.match(r'^[A-Z][^.!?]*$', text))):
                return ContentType.HEADING
        
        # Code block detection
        if self._is_code_like(text) or 'code' in style_name:
            return ContentType.CODE_BLOCK
        
        # List detection
        if (re.match(r'^\s*[-•*]\s+', text) or 
            re.match(r'^\s*\d+\.?\s+', text) or
            'list' in style_name):
            return ContentType.LIST
        
        # Link detection
        if 'http://' in text or 'https://' in text or '@' in text:
            return ContentType.LINK
        
        # Default to text
        return ContentType.TEXT
    
    def _is_code_like(self, text: str) -> bool:
        """Check if text looks like code"""
        code_indicators = [
            '{', '}', '()', '[]', ';', '//', '/*', '*/', 
            'function', 'class', 'def ', 'var ', 'let ', 'const ',
            'import ', 'from ', 'include', '#include'
        ]
        
        # Check for multiple code indicators
        indicator_count = sum(1 for indicator in code_indicators if indicator in text)
        
        # Also check for consistent indentation
        lines = text.split('\n')
        if len(lines) > 1:
            indented_lines = sum(1 for line in lines if line.startswith('    ') or line.startswith('\t'))
            if indented_lines > len(lines) * 0.5:
                indicator_count += 1
        
        return indicator_count >= 2
    
    async def _extract_tables(self,
                             docx_doc: 'DocxDocument',
                             document: ProcessedDocument) -> None:
        """Extract tables from DOCX document"""
        
        try:
            for table_index, table in enumerate(docx_doc.tables):
                # Extract table data
                table_data = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        row_data.append(cell_text)
                    table_data.append(row_data)
                
                if table_data:
                    # Convert table to text representation
                    table_text = self._format_table_as_text(table_data)
                    
                    # Create structured content for table
                    table_content = StructuredContent(
                        content_type=ContentType.TABLE,
                        text=table_text,
                        position={
                            'table_index': table_index
                        },
                        attributes={
                            'rows': len(table_data),
                            'columns': len(table_data[0]) if table_data else 0,
                            'raw_data': table_data
                        },
                        confidence=0.9
                    )
                    
                    document.structured_content.append(table_content)
        
        except Exception as e:
            document.warnings.append(f"Error extracting tables: {str(e)}")
    
    def _format_table_as_text(self, table_data: List[List[str]]) -> str:
        """Format table data as readable text"""
        if not table_data:
            return ""
        
        # Calculate column widths
        col_widths = []
        for col_index in range(len(table_data[0])):
            max_width = max(
                len(str(row[col_index]) if col_index < len(row) else "")
                for row in table_data
            )
            col_widths.append(max(max_width, 3))  # Minimum width of 3
        
        # Format table
        formatted_rows = []
        for row in table_data:
            formatted_cells = []
            for col_index, cell in enumerate(row):
                if col_index < len(col_widths):
                    cell_str = str(cell or "").ljust(col_widths[col_index])
                    formatted_cells.append(cell_str)
            formatted_rows.append(" | ".join(formatted_cells))
        
        return "\n".join(formatted_rows)
    
    def _calculate_confidence_score(self, document: ProcessedDocument) -> float:
        """Calculate confidence score for DOCX processing"""
        
        if document.processing_status == ProcessingStatus.FAILED:
            return 0.0
        
        confidence_factors = []
        
        # Text extraction confidence
        if document.extracted_text:
            text_length = len(document.extracted_text.strip())
            if text_length > 100:
                confidence_factors.append(0.95)  # DOCX is very reliable
            elif text_length > 10:
                confidence_factors.append(0.8)
            else:
                confidence_factors.append(0.4)
        
        # Structured content confidence
        if document.structured_content:
            avg_content_confidence = sum(
                content.confidence for content in document.structured_content
            ) / len(document.structured_content)
            confidence_factors.append(avg_content_confidence)
        
        # Metadata confidence
        if document.metadata and document.metadata.word_count:
            confidence_factors.append(0.9)
        
        # Error penalty
        error_penalty = min(len(document.errors) * 0.1, 0.5)
        
        # Calculate final confidence
        if confidence_factors:
            base_confidence = sum(confidence_factors) / len(confidence_factors)
            return max(0.0, base_confidence - error_penalty)
        else:
            return 0.1  # Minimal confidence if no content extracted
    
    async def extract_images_info(self, file_path: Path) -> List[Dict[str, Any]]:
        """Extract information about images in the DOCX document"""
        
        images_info = []
        
        try:
            docx_doc = DocxDocument(str(file_path))
            
            # Get document part relationships
            doc_part = docx_doc.part
            
            # Find image relationships
            for rel in doc_part.rels.values():
                if "image" in rel.target_ref:
                    image_info = {
                        'relationship_id': rel.rId,
                        'target': rel.target_ref,
                        'content_type': getattr(rel.target_part, 'content_type', 'unknown')
                    }
                    images_info.append(image_info)
            
        except Exception as e:
            raise Exception(f"Error extracting image information: {str(e)}")
        
        return images_info
    
    async def get_document_structure(self, file_path: Path) -> Dict[str, Any]:
        """Get document structure information"""
        
        structure = {
            'paragraphs': 0,
            'tables': 0,
            'images': 0,
            'headings': [],
            'styles_used': set()
        }
        
        try:
            docx_doc = DocxDocument(str(file_path))
            
            # Count paragraphs and analyze headings
            for paragraph in docx_doc.paragraphs:
                if paragraph.text.strip():
                    structure['paragraphs'] += 1
                    
                    # Track styles
                    if paragraph.style:
                        structure['styles_used'].add(paragraph.style.name)
                        
                        # Check for headings
                        if 'heading' in paragraph.style.name.lower():
                            structure['headings'].append({
                                'text': paragraph.text.strip()[:100],  # First 100 chars
                                'style': paragraph.style.name
                            })
            
            # Count tables
            structure['tables'] = len(docx_doc.tables)
            
            # Count images (approximate)
            structure['images'] = len(await self.extract_images_info(file_path))
            
            # Convert set to list for JSON serialization
            structure['styles_used'] = list(structure['styles_used'])
            
        except Exception as e:
            raise Exception(f"Error analyzing document structure: {str(e)}")
        
        return structure